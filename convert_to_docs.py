#!/usr/bin/env python3
"""
마크다운을 DOCX로 변환하는 명령줄 도구
원격 이미지 URL을 포함한 마크다운을 DOCX 파일로 변환합니다.
"""

import argparse
import re
import requests
import io
import sys
from pathlib import Path
from typing import List, Tuple, Optional
import logging

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Error: python-docx가 설치되지 않았습니다.")
    print("다음 명령어로 설치하세요: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
except ImportError:
    print("Error: Pillow가 설치되지 않았습니다.")
    print("다음 명령어로 설치하세요: pip install Pillow")
    sys.exit(1)

# 로깅 설정
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class MarkdownToDocxConverter:
    """마크다운을 DOCX로 변환하는 클래스"""
    
    def __init__(self):
        self.doc = Document()
        self.image_counter = 0
        
    def download_image(self, image_url: str, timeout: int = 30) -> Optional[bytes]:
        """원격 이미지를 다운로드합니다."""
        try:
            logger.info(f"이미지 다운로드 중: {image_url}")
            
            # User-Agent 헤더 추가 (일부 서버에서 요구)
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            response = requests.get(image_url, timeout=timeout, headers=headers)
            response.raise_for_status()
            
            # Content-Type 헤더 확인
            content_type = response.headers.get('Content-Type', '')
            if not content_type.startswith('image/'):
                logger.warning(f"이미지 Content-Type이 아님: {image_url} (Content-Type: {content_type})")
                return None
            
            image_data = response.content
            
            # 이미지 크기 확인
            if len(image_data) > 20 * 1024 * 1024:  # 20MB 제한
                logger.warning(f"이미지가 너무 큽니다 (>{20*1024*1024} bytes): {image_url}")
                return None
            
            if not image_data:
                logger.warning(f"빈 이미지 데이터: {image_url}")
                return None
                
            logger.info(f"이미지 다운로드 완료: {len(image_data)} bytes")
            return image_data
            
        except requests.exceptions.RequestException as e:
            logger.error(f"이미지 다운로드 실패: {image_url} - {str(e)}")
            return None
        except Exception as e:
            logger.error(f"이미지 처리 중 오류: {image_url} - {str(e)}")
            return None
    
    def add_image_to_doc(self, image_data: bytes, alt_text: str = "") -> bool:
        """이미지를 DOCX 문서에 추가합니다."""


        logger.debug(f"add_image_to_doc 호출 중: {alt_text}, {image_data}")
        logger.debug(f"image_data 크기: {len(image_data)} bytes")
        logger.debug(f"image_data 타입: {type(image_data)}")
        try:
            # 이미지 크기 조정 및 검증
            try:
                # DOCX 삽입을 위해 항상 새로운 스트림 생성
                image_stream_for_docx_insertion = io.BytesIO(image_data)
                
                with Image.open(image_stream_for_docx_insertion) as img:
                    logger.info(f"DOCX 삽입용 원본 이미지 형식: {img.format}, 크기: {img.size}")
                    
                    # 이미지 모드 확인 및 변환 (PNG 저장을 위해)
                    if img.mode == 'RGBA':
                        # RGBA를 RGB로 변환 (흰색 배경에 합성)
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        background.paste(img, mask=img.split()[-1])  # 알파 채널을 마스크로 사용
                        img = background
                        logger.info(f"DOCX 삽입용 이미지 모드 변환 완료: RGBA -> RGB (흰색 배경)")
                    elif img.mode != 'RGB':
                        img = img.convert('RGB')
                        logger.info(f"DOCX 삽입용 이미지 모드 변환 완료: {img.mode} -> RGB")
                    
                    width, height = img.size
                    logger.info(f"최종 DOCX 삽입용 이미지 크기: {width}x{height} pixels, 모드: {img.mode}")
                    
                    # 최대 너비 6인치로 제한 (DPI 96 기준)
                    max_width_inches = 6.0
                    dpi = 96  # 기본 DPI
                    
                    # 픽셀을 인치로 변환
                    width_inches = width / dpi
                    height_inches = height / dpi
                    
                    if width_inches > max_width_inches:
                        # 비율 유지하면서 크기 조정
                        scale_factor = max_width_inches / width_inches
                        new_width_inches = max_width_inches
                        new_height_inches = height_inches * scale_factor
                    else:
                        new_width_inches = width_inches
                        new_height_inches = height_inches
                    
                    # 최대 높이도 8인치로 제한
                    max_height_inches = 8.0
                    if new_height_inches > max_height_inches:
                        scale_factor = max_height_inches / new_height_inches
                        new_height_inches = max_height_inches
                        new_width_inches = new_width_inches * scale_factor
                    
                    logger.info(f"조정된 DOCX 삽입용 이미지 크기: {new_width_inches:.2f}x{new_height_inches:.2f} inches")
                    
                    # 최종 이미지를 PNG 형식으로 변환하여 새로운 스트림에 저장
                    final_image_stream = io.BytesIO()
                    img.save(final_image_stream, format='PNG')
                    final_image_stream.seek(0) # 스트림 커서를 처음으로 되돌림
                    
            except Exception as e:
                logger.error(f"DOCX 삽입용 이미지 처리 실패: {str(e)}")
                import traceback
                print(traceback.format_exc())
                return False
                
            # DOCX에 이미지 추가
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 빈 run 생성
            run = paragraph.add_run()
            
            # 최종 이미지 스트림 크기 로깅
            logger.info(f"DOCX에 추가할 PNG 이미지 스트림 크기: {final_image_stream.getbuffer().nbytes} bytes")
            
            run.add_picture(final_image_stream, width=Inches(new_width_inches))
            
            # 이미지 설명 추가 (alt_text가 있는 경우)
            if alt_text:
                caption_paragraph = self.doc.add_paragraph()
                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_paragraph.add_run(f"그림 {self.image_counter + 1}: {alt_text}")
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True
            else:
                # alt_text가 없어도 그림 번호 추가
                caption_paragraph = self.doc.add_paragraph()
                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_paragraph.add_run(f"그림 {self.image_counter + 1}")
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True
            
            self.image_counter += 1
            logger.info(f"DOCX에 이미지 추가 완료: {alt_text or '이미지'}")
            return True
            
        except Exception as e:
            logger.error(f"이미지 추가 실패: {str(e)}")
            return False
    
    def parse_markdown_images(self, markdown_text: str) -> List[Tuple[str, str]]:
        """마크다운에서 이미지 URL을 추출합니다."""
        # 마크다운 이미지 패턴: ![alt text](url)
        image_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        images = re.findall(image_pattern, markdown_text)
        return images
    
    def convert_markdown_to_docx(self, markdown_text: str, output_path: str | None, save_images_to_disk: bool) -> Optional[io.BytesIO]:
        """마크다운을 DOCX로 변환합니다.
        output_path가 제공되면 파일로 저장하고, 그렇지 않으면 BytesIO 객체를 반환합니다.
        """
        try:
            # 이미지 URL 추출
            images = self.parse_markdown_images(markdown_text)
            logger.info(f"발견된 이미지 개수: {len(images)}")
            
            # 모든 이미지를 미리 다운로드
            downloaded_images = {}
            for alt_text, image_url in images:
                logger.info(f"이미지 다운로드 시작: {alt_text} - {image_url}")
                image_data = self.download_image(image_url)
                if image_data:
                    downloaded_images[image_url] = (image_data, alt_text)
                    logger.info(f"이미지 다운로드 성공: {alt_text}")
                else:
                    logger.warning(f"이미지 다운로드 실패: {alt_text} - {image_url}")
            
            logger.info(f"총 {len(downloaded_images)}개 이미지 다운로드 완료")
            
            # 다운로드된 이미지 파일을 로컬 디스크에 저장 (선택적)
            if save_images_to_disk:
                image_id = 0
                image_name_list = []
                import os
                os.makedirs("downloaded_images", exist_ok=True)
                
                for image_url, (image_data, alt_text) in downloaded_images.items():
                    try:
                        # 이미지 이름 생성
                        image_name = f"{image_id}"
                        image_id += 1
                        
                        # 먼저 이미지 형식 확인을 위해 스트림 생성
                        image_stream_for_format_check = io.BytesIO(image_data)
                        
                        file_extension = 'bin'  # 기본 확장자
                        
                        # 이미지 형식 확인
                        try:
                            with Image.open(image_stream_for_format_check) as img:
                                original_format = img.format
                                logger.info(f"저장할 이미지 형식: {original_format}, 크기: {img.size}")
                                
                                # 원본 형식에 따라 확장자 결정
                                if original_format in ['JPEG', 'JPG']:
                                    file_extension = 'jpg'
                                elif original_format == 'PNG':
                                    file_extension = 'png'
                                elif original_format == 'GIF':
                                    file_extension = 'gif'
                                elif original_format == 'BMP':
                                    file_extension = 'bmp'
                                elif original_format == 'TIFF':
                                    file_extension = 'tiff'
                                elif original_format == 'WEBP':
                                    file_extension = 'webp'
                                # 지원하지 않는 형식은 기본값 'bin' 유지
                                
                        except Exception as format_error:
                            # stack trace
                            import traceback
                            print(traceback.format_exc())
                            logger.warning(f"이미지 형식 확인 실패: {str(format_error)}, 기본 확장자 사용")
                            file_extension = 'bin'
                            
                            # Image.open 실패 시 임시 파일로 저장하여 수동 검증
                            try:
                                temp_filename = f"downloaded_images/{image_name}_unidentified.bin"
                                with open(temp_filename, "wb") as f:
                                    f.write(image_data)
                                logger.error(f"식별 실패 이미지 임시 저장: {temp_filename} (수동 확인 필요)")
                            except Exception as e_temp:
                                logger.error(f"임시 파일 저장 실패: {str(e_temp)}")
                        
                        # 원본 데이터를 그대로 저장 (가장 안전한 방법)
                        filename = f"downloaded_images/{image_name}.{file_extension}"
                        with open(filename, "wb") as f:
                            f.write(image_data)
                        image_name_list.append(filename)
                        logger.info(f"이미지 저장 완료: {filename} (원본 데이터)")
                            
                    except Exception as e:
                        logger.error(f"이미지 저장 실패: {alt_text} - {str(e)}")
                        # 실패한 경우에도 시도
                        try:
                            image_name = f"{image_id}_error"
                            image_id += 1
                            filename = f"downloaded_images/{image_name}.bin"
                            with open(filename, "wb") as f:
                                f.write(image_data)
                            image_name_list.append(filename)
                            logger.warning(f"에러 파일로 저장: {filename}")
                        except Exception as e2:
                            logger.error(f"에러 파일 저장도 실패: {str(e2)}")
                
                logger.info(f"총 {len(downloaded_images)}개 이미지 다운로드 완료, {len(image_name_list)}개 이미지 파일 저장 완료")
            else:
                logger.info("다운로드된 이미지 파일을 디스크에 저장하지 않습니다.")

            # 파일을 한 줄씩 읽어서 처리
            lines = markdown_text.split('\n') # 기존 로직 유지, 중복 제거
            current_paragraph = None
            
            for line in lines:
                line = line.strip()
                logger.debug(f"처리 중인 라인: \"{line}\"")
                
                # 이미지 라인 처리 (re.search 사용)
                image_match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', line)
                
                if image_match:
                    logger.debug(f"이미지 패턴 감지됨 (re.search): {image_match.groups()}")
                    alt_text, image_url = image_match.groups()
                    
                    # 미리 다운로드된 이미지 사용
                    if image_url in downloaded_images:
                        image_data, original_alt_text = downloaded_images[image_url]
                        # alt_text가 비어있으면 원본 alt_text 사용
                        display_alt_text = alt_text if alt_text else original_alt_text
                        logger.debug(f"add_image_to_doc 호출 중: {display_alt_text}, {image_url}")
                        success = self.add_image_to_doc(image_data, display_alt_text)
                        if not success:
                            # 이미지 추가 실패 시 플레이스홀더
                            placeholder_paragraph = self.doc.add_paragraph()
                            placeholder_paragraph.add_run(f"[이미지 추가 실패: {display_alt_text}]")
                            placeholder_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            logger.warning(f"DOCX에 이미지 추가 실패: {display_alt_text}")
                    else:
                        # 이미지 다운로드 실패 시 플레이스홀더 추가
                        placeholder_paragraph = self.doc.add_paragraph()
                        placeholder_paragraph.add_run(f"[이미지 로드 실패 (다운로드되지 않음): {image_url}]")
                        placeholder_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        placeholder_paragraph.runs[0].font.color.rgb = None  # 빨간색으로 표시
                        logger.warning(f"DOCX에 이미지 로드 실패 (다운로드되지 않음): {image_url}")
                
                # 헤딩 처리
                elif line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('# ').strip()
                    
                    if level == 1:
                        heading = self.doc.add_heading(text, level=1)
                    elif level == 2:
                        heading = self.doc.add_heading(text, level=2)
                    elif level == 3:
                        heading = self.doc.add_heading(text, level=3)
                    else:
                        heading = self.doc.add_heading(text, level=4)
                
                # 리스트 처리
                elif line.startswith('- ') or line.startswith('* '):
                    if current_paragraph is None:
                        current_paragraph = self.doc.add_paragraph()
                    current_paragraph.add_run(line[2:].strip() + '\n')
                
                # 번호 리스트 처리
                elif re.match(r'^\d+\.\s', line):
                    if current_paragraph is None:
                        current_paragraph = self.doc.add_paragraph()
                    current_paragraph.add_run(line + '\n')
                
                # 인용문 처리
                elif line.startswith('> '):
                    if current_paragraph is None:
                        current_paragraph = self.doc.add_paragraph()
                    current_paragraph.add_run(line[2:].strip() + '\n')
                    current_paragraph.style = 'Quote'
                
                # 코드 블록 처리
                elif line.startswith('```'):
                    if current_paragraph is None:
                        current_paragraph = self.doc.add_paragraph()
                    current_paragraph.add_run(line + '\n')
                    current_paragraph.style = 'No Spacing'
                
                # 일반 텍스트 처리
                elif line:
                    if current_paragraph is None:
                        current_paragraph = self.doc.add_paragraph()
                    
                    # 인라인 포맷팅 처리
                    text = self.process_inline_formatting(line)
                    current_paragraph.add_run(text + '\n')
                
                # 빈 줄 처리
                else:
                    if current_paragraph is not None:
                        current_paragraph = None
            
            # DOCX 파일 저장 또는 BytesIO 반환
            if output_path:
                self.doc.save(output_path)
                logger.info(f"DOCX 파일 저장 완료: {output_path}")
                return io.BytesIO(b'DOCX saved to file') # 성공적으로 파일로 저장되었음을 알리는 더미 BytesIO 반환
            else:
                doc_buffer = io.BytesIO()
                self.doc.save(doc_buffer)
                doc_buffer.seek(0)
                logger.info("DOCX 데이터를 BytesIO 객체로 반환합니다.")
                return doc_buffer
            
        except Exception as e:
            logger.error(f"변환 중 오류 발생: {str(e)}")
            return None
    
    def process_inline_formatting(self, text: str) -> str:
        """인라인 마크다운 포맷팅을 처리합니다."""
        # **bold** 처리
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        # *italic* 처리
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        # `code` 처리
        text = re.sub(r'`(.*?)`', r'\1', text)
        # [link](url) 처리
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        
        return text


def main():
    """메인 함수"""
    parser = argparse.ArgumentParser(description='마크다운을 DOCX로 변환합니다.')
    parser.add_argument('input_file', help='입력 마크다운 파일 경로')
    parser.add_argument('-o', '--output', help='출력 DOCX 파일 경로 (기본값: 입력파일명.docx)')
    parser.add_argument('-v', '--verbose', action='store_true', help='상세 로그 출력')
    parser.add_argument('--save-images-to-disk', action='store_true', help='다운로드된 이미지 파일을 로컬 디스크에 저장합니다 (기본값: 저장 안 함)')
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # 입력 파일 확인
    input_path = Path(args.input_file)
    if not input_path.exists():
        logger.error(f"입력 파일을 찾을 수 없습니다: {input_path}")
        sys.exit(1)
    
    # 출력 파일 경로 설정
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.with_suffix('.docx')
    
    # 마크다운 파일 읽기
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
    except Exception as e:
        logger.error(f"파일 읽기 실패: {str(e)}")
        sys.exit(1)
    
    # 변환 실행
    converter = MarkdownToDocxConverter()
    success = converter.convert_markdown_to_docx(markdown_text, str(output_path), args.save_images_to_disk)
    
    if success:
        logger.info(f"변환 완료: {output_path}")
        print(f"✅ 변환 완료: {output_path}")
    else:
        logger.error("변환 실패")
        print("❌ 변환 실패")
        sys.exit(1)


if __name__ == '__main__':
    main()
